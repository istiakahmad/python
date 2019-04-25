#pip install python-docx
import docx

#Writing Word Documents
doc = docx.Document()

# Writting Heading
doc.add_heading('Header 0 yo yo', 0)

#Writing Paragraph
doc.add_paragraph('Hello world!')
doc.add_heading('Header 1 ho ho ho', 4)
paraObj1 = doc.add_paragraph('This is a second paragraph.')
doc.paragraphs[3].runs[0].add_break() # Line Break after 3 paragraph

paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph.')

#Add new page
doc.paragraphs[4].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE) # After 4 paragraph page break
doc.add_paragraph('This is on the second page!')

#Adding Picture with width and height
doc.add_picture('G:\i.jpg', width=docx.shared.Inches(1), height=docx.shared.Cm(4))
#doc.add_picture('G:\i.jpg') # image with original size

doc.save('D:/python/helloworld.docx')


#Read Document Text

doc = docx.Document('D:/python/doc.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(doc.paragraphs[4].runs[0].text)
print(doc.paragraphs[4].runs[1].text)
print(doc.paragraphs[4].runs[2].text)


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('D:/python/doc.docx'))
